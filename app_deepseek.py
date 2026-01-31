import streamlit as st
from openai import OpenAI
import difflib
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image
import pytesseract
import os
import shutil

# --- 0. Tesseract 路径强制修复 ---
if os.path.exists('/usr/bin/tesseract'):
    pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
else:
    possible_path = shutil.which("tesseract")
    if possible_path:
        pytesseract.pytesseract.tesseract_cmd = possible_path

# --- 1. 页面配置 ---
st.set_page_config(
    page_title="Ketty's Mini Proofreading", 
    page_icon="✒️", 
    layout="centered"
)

# --- 2. 状态初始化 ---
# 这里初始化默认模式，如果用户没选过，默认是 "仅标红"
if "selected_mode" not in st.session_state:
    st.session_state.selected_mode = "仅标红"

if 'ocr_text' not in st.session_state:
    st.session_state['ocr_text'] = ""

# --- 3. CSS 样式 (仅保留基础美化，移除了丑陋的Radio样式) ---
def local_css():
    st.markdown("""
    <style>
    .stApp {
        background-color: #ffffff;
        font-family: "PingFang SC", "Microsoft YaHei", -apple-system, sans-serif;
    }
    .nav-title {
        font-size: 22px;
        font-weight: 700;
        color: #1a1a1a;
        margin-bottom: 20px;
        text-align: center;
    }
    /* 描述文字样式 */
    .mode-desc {
        background-color: #f9fafb;
        padding: 15px;
        border-radius: 8px;
        border-left: 4px solid #1a1a1a;
        color: #4b5563;
        font-size: 14px;
        margin: 15px 0;
        line-height: 1.6;
    }
    /* 输入框样式 */
    .stTextArea textarea {
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 16px;
        background-color: #fcfcfc;
    }
    .stTextArea textarea:focus {
        border-color: #1a1a1a;
        box-shadow: 0 0 0 2px rgba(0,0,0,0.1);
    }
    /* 隐藏页脚 */
    #MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)

local_css()

# --- 4. API 初始化 ---
try:
    if "DEEPSEEK_API_KEY" in st.secrets:
        api_key = st.secrets["DEEPSEEK_API_KEY"]
    else:
        st.stop()
except:
    st.stop()

client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

# --- 5. 顶部导航栏 (交互式按钮组) ---
st.markdown('<div class="nav-title">✒️ Ketty\'s Mini Proofreading</div>', unsafe_allow_html=True)

# 定义回调函数：点击按钮时更新 session_state
def set_mode(mode):
    st.session_state.selected_mode = mode

# 创建三列布局放置按钮
col_nav1, col_nav2, col_nav3 = st.columns(3)

with col_nav1:
    # 如果当前模式是"仅标红"，按钮显示为实心(primary)，否则为轮廓(secondary)
    st.button(
        "仅标红", 
        type="primary" if st.session_state.selected_mode == "仅标红" else "secondary", 
        use_container_width=True,
        on_click=set_mode,
        args=("仅标红",)
    )

with col_nav2:
    st.button(
        "纠错", 
        type="primary" if st.session_state.selected_mode == "纠错" else "secondary", 
        use_container_width=True,
        on_click=set_mode,
        args=("纠错",)
    )

with col_nav3:
    st.button(
        "润色", 
        type="primary" if st.session_state.selected_mode == "润色" else "secondary", 
        use_container_width=True,
        on_click=set_mode,
        args=("润色",)
    )

st.markdown("---")

# --- 6. 动态内容配置 ---
# 获取当前选中的模式
current_mode = st.session_state.selected_mode

mode_config = {
    "仅标红": {
        "desc": "🔴 **Strict Mode**：严格查错，仅标红原文中的错别字与语病，**绝不改写**。",
        "placeholder": "请输入原文...",
        "btn_text": "开始扫描 / Strict Scan",
        "prompt": """
            你是一个严格的校对员。请检查文本中的【错别字】、【标点错误】和【明显语病】。
            【绝对指令】：
            1. 严禁重写句子，严禁润色，严禁改变原意。
            2. 【重要】输出文本必须与原文段落结构、换行符、字数行数高度一致。严禁合并段落。
            3. 如果没有错误，请原样输出。
            直接输出修正后的全文，不含解释。
        """
    },
    "纠错": {
        "desc": "🛠️ **Fix Mode**：智能修正错别字、标点及不通顺语句，保持原意。",
        "placeholder": "请输入原文...",
        "btn_text": "开始纠错 / Auto Fix",
        "prompt": """
            你是一个语文老师。修正错别字、语病和标点。
            【重要指令】：
            1. 保持原文语气，只确保规范。
            2. 【严禁合并段落】：必须严格保留原文的换行符和段落结构，原文有几段，输出就是几段。
            直接输出修正后的文本，不要加任何前言后语。
        """
    },
    "润色": {
        "desc": "✨ **Polish Mode**：深度优化用词与句式，提升文章的专业度与文采。",
        "placeholder": "请输入原文...",
        "btn_text": "开始润色 / Polish Magic",
        "prompt": """
            你是一个资深的编辑。请对文本进行深度润色，优化用词和句式，使其更加流畅专业。
            【重要指令】：
            1. 提升文采，但不要过度改变原意。
            2. 【严禁合并段落】：输出必须严格保留原文的段落结构和换行，不要将文本合并成一大段。
            直接输出结果，不要加任何解释。
        """
    }
}

current_config = mode_config[current_mode]
st.markdown(f'<div class="mode-desc">{current_config["desc"]}</div>', unsafe_allow_html=True)

# --- 7. 图片上传功能 ---
with st.expander("🖼️ 上传图片识别文字 / Upload Image OCR"):
    uploaded_file = st.file_uploader("选择一张图片 (支持 JPG/PNG)", type=['png', 'jpg', 'jpeg'])
    
    if uploaded_file is not None:
        try:
            with st.spinner("正在识别图片文字..."):
                img = Image.open(uploaded_file)
                text_from_image = pytesseract.image_to_string(img, lang='chi_sim+eng')
                
                if text_from_image.strip():
                    st.session_state['ocr_text'] = text_from_image.strip()
                    st.success("✅ 识别成功！文字已填入下方输入框。")
                    st.rerun() # 识别后刷新页面以填入文字
                else:
                    st.warning("⚠️ 图片中未识别到清晰文字。")
        except pytesseract.TesseractNotFoundError:
            st.error("❌ 核心错误：云端服务器未安装 Tesseract 引擎。")
        except Exception as e:
            st.error(f"识别出错: {e}")

# --- 8. 输入区 ---
final_value = st.session_state['ocr_text'] if st.session_state['ocr_text'] else ""

text_input = st.text_area(
    "",
    height=300,
    placeholder=current_config["placeholder"],
    value=final_value, 
    key="main_input"
)

# 按钮
run_btn = st.button(current_config["btn_text"], type="primary")

# --- 9. 执行逻辑 ---
if run_btn:
    if not text_input:
        st.warning("⚠️ 请先输入文字内容")
    else:
        with st.spinner("DeepSeek AI 正在思考中..."):
            try:
                response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {"role": "system", "content": current_config["prompt"]},
                        {"role": "user", "content": text_input}
                    ],
                    stream=False
                )
                res_text = response.choices[0].message.content.strip()

                # 结果展示样式
                st.markdown(
                    """
                    <style>
                    .result-box {
                        margin-top: 25px;
                        padding: 40px;
                        border: 2px dashed #e5e7eb;
                        border-radius: 4px;
                        background: #ffffff;
                        font-family: "Songti SC", "SimSun", serif; 
                        font-size: 18px;
                        line-height: 2.0;
                        white-space: pre-wrap;
                        word-wrap: break-word;
                    }
                    </style>
                    """, unsafe_allow_html=True
                )

                def get_diff_html(orig, corr, mode):
                    output = []
                    s = difflib.SequenceMatcher(None, orig, corr, autojunk=False)
                    for opcode, a0, a1, b0, b1 in s.get_opcodes():
                        if mode == "仅标红":
                            if opcode == 'equal':
                                output.append(f'<span>{orig[a0:a1]}</span>')
                            elif opcode in ['delete', 'replace']:
                                output.append(f'<span style="color:#e11d48; font-weight:bold; background-color:#fff1f2; padding:0 2px;">{orig[a0:a1]}</span>')
                            elif opcode == 'insert':
                                output.append(f'<span style="color:#e11d48; font-weight:bold;">^</span>')
                        else:
                            if opcode == 'equal':
                                output.append(orig[a0:a1])
                            elif opcode == 'insert':
                                output.append(f'<span style="color:#059669; font-weight:bold;">{corr[b0:b1]}</span>')
                            elif opcode in ['delete', 'replace']:
                                output.append(f'<span style="color:#9ca3af; text-decoration:line-through;">{orig[a0:a1]}</span>')
                                if opcode == 'replace':
                                    output.append(f'<span style="color:#059669; font-weight:bold;">{corr[b0:b1]}</span>')
                    return "".join(output)

                html_content = get_diff_html(text_input, res_text, current_mode)
                st.markdown(f'<div class="result-box">{html_content}</div>', unsafe_allow_html=True)
                
                # Word 导出逻辑
                def create_docx(orig, corr, mode):
                    doc = Document()
                    doc.add_heading(f'Ketty\'s Review - {mode}', 0)
                    style = doc.styles['Normal']
                    style.font.name = 'SimSun'
                    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    p = doc.add_paragraph()
                    s = difflib.SequenceMatcher(None, orig, corr, autojunk=False)
                    for opcode, a0, a1, b0, b1 in s.get_opcodes():
                        if mode == "仅标红":
                            if opcode == 'equal':
                                run = p.add_run(orig[a0:a1])
                                run.font.color.rgb = RGBColor(0,0,0)
                            elif opcode in ['delete', 'replace']:
                                run = p.add_run(orig[a0:a1])
                                run.font.color.rgb = RGBColor(255,0,0)
                            elif opcode == 'insert':
                                run = p.add_run("^")
                                run.font.color.rgb = RGBColor(255,0,0)
                                run.font.bold = True
                        else:
                            p.add_run(corr)
                    f = BytesIO()
                    doc.save(f)
                    f.seek(0)
                    return f

                st.markdown("<br>", unsafe_allow_html=True)
                file_docx = create_docx(text_input, res_text, current_mode)
                st.download_button(
                    label=f"📥 导出报告 / Download (.docx)",
                    data=file_docx,
                    file_name=f"Ketty_{current_mode}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"Error: {e}")
